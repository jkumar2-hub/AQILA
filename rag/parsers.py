"""
AQILA — Document Parser Module
Owner: M1 — AI / ML + RAG Core

Responsibilities:
    - PDF parsing
        Primary: Docling
        Fallback: PyMuPDF
    - DOCX parsing
        python-docx
    - Text chunking
        400-token target
        50-token overlap

Output:
    list[dict] containing chunk text and metadata.
"""

from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# PDF PARSING
# ---------------------------------------------------------------------------

def parse_pdf(file_path: str | Path) -> list[dict[str, Any]]:
    """
    Parse a PDF into page-level text records.

    Docling is attempted first.
    PyMuPDF is used as the fallback if Docling fails.

    Returns:
        [
            {
                "page_number": int,
                "text": str,
            },
            ...
        ]
    """

    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"PDF not found: {file_path}")

    # ---------------------------------------------------------
    # Primary parser: Docling
    # ---------------------------------------------------------

    try:
        from docling.document_converter import DocumentConverter

        converter = DocumentConverter()
        result = converter.convert(str(file_path))

        document = result.document

        # Export page-aware markdown/text.
        #
        # Docling versions differ in their page APIs, so we
        # gracefully fall back to document-level text if needed.
        try:
            pages = []

            for page_number, page in enumerate(
                document.pages.values(),
                start=1,
            ):
                text = ""

                if hasattr(page, "export_to_markdown"):
                    text = page.export_to_markdown()

                elif hasattr(page, "export_to_text"):
                    text = page.export_to_text()

                if text and text.strip():
                    pages.append(
                        {
                            "page_number": page_number,
                            "text": text.strip(),
                        }
                    )

            if pages:
                return pages

        except Exception:
            # Docling successfully parsed the document but
            # page-level extraction was unavailable.
            pass

        # Document-level Docling fallback.
        try:
            text = document.export_to_markdown()

            if text and text.strip():
                return [
                    {
                        "page_number": 1,
                        "text": text.strip(),
                    }
                ]
        except Exception:
            pass

    except Exception as exc:
        print(
            f"[AQILA] Docling PDF parsing failed: {exc}"
        )

    # ---------------------------------------------------------
    # Fallback parser: PyMuPDF
    # ---------------------------------------------------------

    return _parse_pdf_pymupdf(file_path)


def _parse_pdf_pymupdf(
    file_path: Path,
) -> list[dict[str, Any]]:
    """
    Parse PDF using PyMuPDF.

    This is the immediate fallback when Docling fails.
    """

    import fitz

    pages: list[dict[str, Any]] = []

    with fitz.open(file_path) as document:
        for page_index, page in enumerate(document):
            text = page.get_text("text").strip()

            if not text:
                continue

            pages.append(
                {
                    "page_number": page_index + 1,
                    "text": text,
                }
            )

    return pages


# ---------------------------------------------------------------------------
# DOCX PARSING
# ---------------------------------------------------------------------------

def parse_docx(file_path: str | Path) -> list[dict[str, Any]]:
    """
    Parse a DOCX document using python-docx.

    Returns:
        [
            {
                "page_number": None,
                "text": str,
            }
        ]
    """

    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(
            f"DOCX not found: {file_path}"
        )

    from docx import Document

    document = Document(file_path)

    paragraphs: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    # Include table content as text too.
    for table in document.tables:
        for row in table.rows:
            cells = []

            for cell in row.cells:
                text = cell.text.strip()

                if text:
                    cells.append(text)

            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n".join(paragraphs).strip()

    if not text:
        return []

    return [
        {
            "page_number": None,
            "text": text,
        }
    ]


# ---------------------------------------------------------------------------
# TOKEN-BASED CHUNKING
# ---------------------------------------------------------------------------

def chunk_text(
    text: str,
    chunk_size: int = 400,
    overlap: int = 50,
) -> list[str]:
    """
    Split text into overlapping chunks.

    Target:
        chunk_size = 400 tokens
        overlap = 50 tokens

    For the initial implementation, whitespace-separated
    words are used as a lightweight token approximation.

    This keeps the parser dependency-light and deterministic.
    """

    if not text or not text.strip():
        return []

    if overlap >= chunk_size:
        raise ValueError(
            "overlap must be smaller than chunk_size"
        )

    words = text.split()

    if not words:
        return []

    chunks: list[str] = []

    step = chunk_size - overlap

    for start in range(0, len(words), step):
        end = start + chunk_size

        chunk = " ".join(words[start:end]).strip()

        if chunk:
            chunks.append(chunk)

        if end >= len(words):
            break

    return chunks


# ---------------------------------------------------------------------------
# SOURCE-AWARE CHUNKING
# ---------------------------------------------------------------------------

def parse_and_chunk(
    file_path: str | Path,
    source_id: str,
    file_name: str,
    modality: str = "text",
) -> list[dict[str, Any]]:
    """
    Parse a supported document and convert it into chunk records.

    Output metadata is designed to feed the M1 embedding/indexing
    pipeline later.
    """

    file_path = Path(file_path)

    extension = file_path.suffix.lower()

    if extension == ".pdf":
        pages = parse_pdf(file_path)

    elif extension == ".docx":
        pages = parse_docx(file_path)

    else:
        raise ValueError(
            f"Unsupported document type: {extension}"
        )

    chunks: list[dict[str, Any]] = []

    chunk_index = 0

    for page in pages:
        page_number = page.get("page_number")
        page_text = page.get("text", "")

        page_chunks = chunk_text(
            page_text,
            chunk_size=400,
            overlap=50,
        )

        for chunk in page_chunks:
            chunks.append(
                {
                    "source_id": source_id,
                    "file_name": file_name,
                    "modality": modality,
                    "page_number": page_number,
                    "chunk_index": chunk_index,
                    "text": chunk,
                }
            )

            chunk_index += 1

    return chunks